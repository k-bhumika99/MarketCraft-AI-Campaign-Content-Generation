"""
kit_export.py — Export & Asset Management Module

Lets users preview, organize, and export the complete marketing kit
generated for a campaign in PDF, DOCX, and TXT formats.
"""
import io

from reportlab.lib.pagesizes import A4
from reportlab.lib.units import mm
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, ListFlowable, ListItem
)

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

BRAND_PRIMARY = colors.HexColor("#6C63FF")
BRAND_ACCENT = colors.HexColor("#FF6FA8")
BRAND_INK = colors.HexColor("#241F3A")
BRAND_MUTED = colors.HexColor("#6E6689")


def _kit_text_blocks(campaign):
    info = campaign.get("campaign_info", {}) or {}
    content = campaign.get("content", {}) or {}
    sm = content.get("social_media", {}) or {}
    mc = content.get("marketing_content", {}) or {}
    seo = content.get("seo", {}) or {}
    cd = content.get("creative_design", {}) or {}
    val = content.get("validation", {}) or {}

    blocks = []
    blocks.append(("Campaign Overview", [
        f"Campaign: {info.get('campaign_name', '')}",
        f"Product: {info.get('product_name', '')}",
        f"Objective: {info.get('campaign_objective', '')}",
        f"Target Audience: {info.get('target_audience', '')}",
        f"Brand Tone: {info.get('brand_tone', '')}",
        f"Key Message: {info.get('key_message', '')}",
    ]))

    blocks.append(("Social Media Content", [
        f"Instagram: {sm.get('instagram', {}).get('caption', '')}",
        f"Instagram hashtags: {' '.join(sm.get('instagram', {}).get('hashtags', []))}",
        f"Facebook: {sm.get('facebook', {}).get('post', '')}",
        f"LinkedIn: {sm.get('linkedin', {}).get('post', '')}",
        f"Twitter/X: {sm.get('twitter_x', {}).get('post', '')}",
        f"Google Ads headlines: {' | '.join(sm.get('google_ads', {}).get('headlines', []))}",
        f"Google Ads descriptions: {' | '.join(sm.get('google_ads', {}).get('descriptions', []))}",
    ]))

    blocks.append(("Marketing Copy", [
        f"Headlines: {' | '.join(mc.get('headlines', []))}",
        f"Taglines: {' | '.join(mc.get('taglines', []))}",
        f"CTAs: {' | '.join(mc.get('ctas', []))}",
        f"Short description: {mc.get('short_description', '')}",
        f"Long description: {mc.get('long_description', '')}",
    ]))

    blocks.append(("SEO", [
        f"Keywords: {', '.join(seo.get('keywords', []))}",
        f"Hashtags: {' '.join(seo.get('hashtags', []))}",
        f"Meta title: {seo.get('meta_title', '')}",
        f"Meta description: {seo.get('meta_description', '')}",
    ]))

    scenes = cd.get("video_ad_script", {}).get("scenes", [])
    scene_lines = [f"Scene {s.get('scene')}: {s.get('visual')} — VO: \"{s.get('voiceover')}\" ({s.get('duration_sec')}s)" for s in scenes]
    blocks.append(("Creative Design", [
        f"Poster concept: {cd.get('poster_concept', '')}",
        f"Banner copy: {cd.get('banner_copy', '')}",
        f"Story creative: {cd.get('story_creative', '')}",
        "Video ad script:",
        *scene_lines,
    ]))

    blocks.append(("Validation", [
        f"Overall score: {val.get('overall_score', '')}/100",
        f"Consistency: {val.get('consistency_score', '')}/100",
        f"Brand alignment: {val.get('brand_alignment_score', '')}/100",
        f"Platform compatibility: {val.get('platform_compatibility_score', '')}/100",
        *[f"Note: {n}" for n in val.get("notes", [])],
    ]))

    return blocks


def campaign_txt(campaign) -> bytes:
    lines = [f"MARKETCRAFT AI — MARKETING KIT", "=" * 40, ""]
    for title, items in _kit_text_blocks(campaign):
        lines.append(title.upper())
        lines.append("-" * len(title))
        lines.extend(items)
        lines.append("")
    return "\n".join(lines).encode("utf-8")


def campaign_pdf(campaign) -> bytes:
    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4, topMargin=18 * mm, bottomMargin=18 * mm,
                             leftMargin=18 * mm, rightMargin=18 * mm)
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle("KitTitle", parent=styles["Title"], textColor=BRAND_PRIMARY, fontSize=22)
    h2 = ParagraphStyle("KitH2", parent=styles["Heading2"], textColor=BRAND_PRIMARY, spaceBefore=14, spaceAfter=6)
    body = ParagraphStyle("KitBody", parent=styles["BodyText"], textColor=BRAND_INK, fontSize=10, leading=14)

    info = campaign.get("campaign_info", {}) or {}
    story = [
        Paragraph("MarketCraft AI", title_style),
        Paragraph(f"Marketing Kit — {info.get('campaign_name', 'Campaign')}", styles["Heading3"]),
        Spacer(1, 10),
    ]

    for title, items in _kit_text_blocks(campaign):
        story.append(Paragraph(title, h2))
        for item in items:
            if item:
                story.append(Paragraph(item.replace("&", "&amp;"), body))
        story.append(Spacer(1, 4))

    doc.build(story)
    return buf.getvalue()


def campaign_docx(campaign) -> bytes:
    d = Document()
    info = campaign.get("campaign_info", {}) or {}

    title = d.add_heading("MarketCraft AI — Marketing Kit", level=0)
    d.add_heading(info.get("campaign_name", "Campaign"), level=1)

    for section_title, items in _kit_text_blocks(campaign):
        d.add_heading(section_title, level=2)
        for item in items:
            if item:
                d.add_paragraph(item)

    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()
