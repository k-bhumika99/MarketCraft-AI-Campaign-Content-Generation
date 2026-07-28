"""
report_parser.py — Campaign Report Import Module

Accepts campaign reports (PDF/DOCX) — typically produced by an upstream
Campaign Planning Agent — extracts the raw text, and validates that there
is enough content for the Campaign Understanding Agent to work with.
"""
import io
import os

import fitz  # PyMuPDF
import docx  # python-docx


class ReportParseError(Exception):
    pass


ALLOWED_EXTENSIONS = {".pdf", ".docx", ".txt"}


def allowed_file(filename: str) -> bool:
    ext = os.path.splitext(filename or "")[1].lower()
    return ext in ALLOWED_EXTENSIONS


def extract_text(file_bytes: bytes, filename: str) -> str:
    """Extract raw text from an uploaded campaign report (PDF/DOCX/TXT)."""
    ext = os.path.splitext(filename or "")[1].lower()

    if ext == ".pdf":
        text = _extract_pdf(file_bytes)
    elif ext == ".docx":
        text = _extract_docx(file_bytes)
    elif ext == ".txt":
        text = file_bytes.decode("utf-8", errors="ignore")
    else:
        raise ReportParseError(
            "Unsupported file type. Please upload a campaign report as PDF, DOCX, or TXT."
        )

    text = (text or "").strip()
    if len(text) < 40:
        raise ReportParseError(
            "Couldn't find enough readable text in that report. "
            "Try a different file, or use manual entry instead."
        )
    return text


def _extract_pdf(file_bytes: bytes) -> str:
    try:
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        parts = [page.get_text() for page in doc]
        doc.close()
        return "\n".join(parts)
    except Exception as exc:
        raise ReportParseError(f"Could not read PDF: {exc}")


def _extract_docx(file_bytes: bytes) -> str:
    try:
        d = docx.Document(io.BytesIO(file_bytes))
        parts = [p.text for p in d.paragraphs]
        for table in d.tables:
            for row in table.rows:
                parts.append(" | ".join(c.text for c in row.cells))
        return "\n".join(parts)
    except Exception as exc:
        raise ReportParseError(f"Could not read DOCX: {exc}")


def validate_data(campaign_info: dict) -> list:
    """Basic validation of the Campaign Understanding Agent's output.
    Returns a list of warning strings (empty if all good)."""
    warnings = []
    required = ["product_name", "target_audience", "campaign_objective"]
    for field in required:
        if not campaign_info.get(field):
            warnings.append(f"Missing '{field.replace('_', ' ')}' — content quality may suffer.")
    return warnings
